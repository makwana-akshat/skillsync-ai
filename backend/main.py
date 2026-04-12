import os
from dotenv import load_dotenv
load_dotenv()

from fastapi import FastAPI, UploadFile, File, Form, Query
from fastapi.responses import FileResponse
from typing import List, Optional
from fastapi.middleware.cors import CORSMiddleware
from parser import parse_resume, extract_skills_from_text, extract_job_skills
from normalizer import normalize_skills
from hierarchy import infer_hierarchy
from skill_registry import process_skills
from matcher import match_skills
from report import generate_pdf
from table_extractor import extract_tables_from_bytes
import asyncio
import json
import hashlib
from pydantic import BaseModel, EmailStr
from database import init_db, add_analysis_record, get_overview_stats, get_all_records, create_user, get_user_by_email
from auth import get_password_hash, verify_password, create_access_token, get_current_user
from fastapi import HTTPException, status, Depends
from fastapi.security import OAuth2PasswordRequestForm

init_db()

app = FastAPI(title="SkillSync AI", description="AI-based system for parsing and matching skills")

# Setup CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── Default tier thresholds ──────────────────────────────────────────
DEFAULT_THRESHOLDS = {"top": 80, "mid": 60, "accept": 60}

def classify_tier(score: float, thresholds: dict = None) -> str:
    """Classify a score into Top / Mid / Low Tier."""
    t = thresholds or DEFAULT_THRESHOLDS
    if score >= t.get("top", 80):
        return "Top Tier"
    elif score >= t.get("mid", 60):
        return "Mid Tier"
    else:
        return "Low Tier"

def classify_decision(score: float, thresholds: dict = None) -> str:
    """Return ACCEPTED / REJECTED based on the acceptance threshold."""
    t = thresholds or DEFAULT_THRESHOLDS
    # Use 'accept' key if present, otherwise fall back to 'mid'
    accept_val = t.get("accept", t.get("mid", 60))
    return "ACCEPTED" if score >= accept_val else "REJECTED"


# ── Helpers for JD file extraction ───────────────────────────────────
def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from uploaded PDF, DOCX, or TXT file for JD."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    
    if ext == "txt":
        try:
            return file_bytes.decode("utf-8", errors="replace")
        except Exception:
            return file_bytes.decode("latin-1", errors="replace")
    
    elif ext == "pdf":
        try:
            import pdfplumber
            import io
            text_parts = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)
            return "\n".join(text_parts) if text_parts else ""
        except Exception as e:
            print(f"PDF JD extraction error: {e}")
            return ""
    
    elif ext in ("docx", "doc"):
        try:
            import docx
            import io
            doc = docx.Document(io.BytesIO(file_bytes))
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception as e:
            print(f"DOCX JD extraction error: {e}")
            return ""
    
    return ""


# ── Duplicate detection helper ───────────────────────────────────────
def deduplicate_files(file_contents: list) -> list:
    """Remove duplicate files based on content hash or filename.
    file_contents: list of (bytes, filename) tuples.
    Returns deduplicated list, keeping only first occurrence.
    """
    seen_hashes = set()
    seen_names = set()
    unique = []
    
    for fbytes, fname in file_contents:
        content_hash = hashlib.md5(fbytes).hexdigest()
        # Skip if same content hash OR same filename
        if content_hash in seen_hashes or fname in seen_names:
            continue
        seen_hashes.add(content_hash)
        seen_names.add(fname)
        unique.append((fbytes, fname))
    
    return unique


# ── Auth models ─────────────────────────────────────────────────────
class UserSignup(BaseModel):
    company_name: str
    email: EmailStr
    password: str

class UserLogin(BaseModel):
    email: EmailStr
    password: str

@app.post("/signup", status_code=status.HTTP_201_CREATED)
def signup(user: UserSignup):
    hashed_password = get_password_hash(user.password)
    success = create_user(user.company_name, user.email, hashed_password)
    if not success:
        raise HTTPException(status_code=400, detail="Email already exists")
    return {"message": "User created successfully"}

@app.post("/login")
def login(user: UserLogin):
    db_user = get_user_by_email(user.email)
    if not db_user or not verify_password(user.password, db_user["hashed_password"]):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid credentials",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token = create_access_token(data={"sub": user.email, "user_id": db_user["id"]})
    return {"access_token": access_token, "token_type": "bearer", "user": {"email": db_user["email"], "company_name": db_user["company_name"]}}


@app.get("/")
def read_root():
    return {"message": "Welcome to SkillSync AI"}

@app.get("/stats")
def stats_endpoint(current_user: dict = Depends(get_current_user)):
    return get_overview_stats(current_user["id"])

@app.get("/history")
def history_endpoint(
    limit: int = Query(200, ge=1, le=500),
    sort: str = Query("date"),
    order: str = Query("desc"),
    tier: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    current_user: dict = Depends(get_current_user),
):
    """Return analysis history records with sorting & filtering."""
    return {"records": get_all_records(current_user["id"], limit, sort, order, tier, status)}

@app.post("/test")
async def test_endpoint(
    file: UploadFile = File(...),
    job_description: str = Form(...)
):
    return {
        "filename": file.filename,
        "job_description": job_description,
        "message": "API working"
    }

@app.post("/parse")
async def parse_endpoint(file: UploadFile = File(...)):
    file_bytes = await file.read()
    extracted_data = parse_resume(file_bytes, file.filename)
    
    if "error" in extracted_data:
        return {"filename": file.filename, "error": extracted_data["error"], "extracted_skills": []}
    
    raw_skills = extracted_data.get("skills", [])
    
    # Process skills: Normalize -> Hierarchy -> Registry
    normalized = normalize_skills(raw_skills)
    final_skills = infer_hierarchy(normalized)
    registry_result = process_skills(final_skills)
    
    return {
        "filename": file.filename,
        "extracted_skills": final_skills,
        "processed_skills": registry_result.get("processed_skills", []),
        "emerging_skills": registry_result.get("emerging_candidates", []),
        "store_updates": registry_result.get("store_updates", []),
        "raw_skills": raw_skills,
    }

@app.post("/analyze")
async def analyze_endpoint(
    file: UploadFile = File(...),
    job_description: str = Form(""),
    thresholds: Optional[str] = Form(None),
    jd_file: Optional[UploadFile] = File(None),
    current_user: dict = Depends(get_current_user)
):
    # 1. Read file bytes
    file_bytes = await file.read()

    # Parse thresholds from frontend (JSON string)
    t = DEFAULT_THRESHOLDS.copy()
    if thresholds:
        try:
            t.update(json.loads(thresholds))
        except (json.JSONDecodeError, TypeError):
            pass
    
    # Handle JD file upload — extract text and merge with text input
    jd_text = job_description or ""
    if jd_file:
        jd_file_bytes = await jd_file.read()
        extracted_jd = extract_text_from_file(jd_file_bytes, jd_file.filename)
        if extracted_jd:
            # Merge: existing text first, then file content
            jd_text = (jd_text.strip() + "\n\n" + extracted_jd.strip()).strip()
    
    if not jd_text.strip():
        return {"error": "Please provide a job description (text or file).", "score": 0}
    
    # Run parsing and fetching job skills concurrently
    extracted_data, job_skills = await asyncio.gather(
        asyncio.to_thread(parse_resume, file_bytes, file.filename),
        asyncio.to_thread(extract_job_skills, jd_text)
    )
    
    if "error" in extracted_data:
        return {"error": extracted_data["error"], "score": 0, "matched_skills": [], "missing_skills": []}
        
    resume_skills = extracted_data.get("skills", [])
    
    # 3. Match skills
    match_result = match_skills(resume_skills, job_skills, jd_text)
    
    # 4. Classify
    score = match_result["score"]
    decision = classify_decision(score, t)
    tier = classify_tier(score, t)
    
    # 5. Save to Database
    add_analysis_record(current_user["id"], file.filename, score, decision, tier)
    
    # 6. Output — include categorized skills for frontend
    return {
        "score": score,
        "decision": decision,
        "tier": tier,
        "thresholds": t,
        "keyword_score": match_result.get("keyword_score", 0.0),
        "semantic_score": match_result.get("semantic_score", 0.0),
        "matched_skills": match_result.get("matched_skills", []),
        "top_semantic_matches": match_result.get("top_semantic_matches", []),
        "matched_required": match_result["matched_required"],
        "matched_optional": match_result["matched_optional"],
        "missing_required": match_result["missing_required"],
        "missing_optional": match_result["missing_optional"],
        "explanation": match_result["explanation"],
        "processed_skills": match_result.get("processed_skills", []),
        "emerging_skills": match_result.get("emerging_skills", [])
    }

@app.post("/rank")
async def rank_endpoint(
    files: List[UploadFile] = File(...),
    job_description: str = Form(""),
    thresholds: Optional[str] = Form(None),
    jd_file: Optional[UploadFile] = File(None),
    current_user: dict = Depends(get_current_user)
):
    # Parse thresholds
    t = DEFAULT_THRESHOLDS.copy()
    if thresholds:
        try:
            t.update(json.loads(thresholds))
        except (json.JSONDecodeError, TypeError):
            pass

    # Handle JD file upload
    jd_text = job_description or ""
    if jd_file:
        jd_file_bytes = await jd_file.read()
        extracted_jd = extract_text_from_file(jd_file_bytes, jd_file.filename)
        if extracted_jd:
            jd_text = (jd_text.strip() + "\n\n" + extracted_jd.strip()).strip()
    
    if not jd_text.strip():
        return {"error": "Please provide a job description (text or file)."}

    # 1. Read all files first
    file_contents = []
    for file in files:
        content = await file.read()
        file_contents.append((content, file.filename))
    
    # 1b. Deduplicate files
    file_contents = deduplicate_files(file_contents)
    
    # 2. Create parallel tasks
    tasks = [asyncio.to_thread(extract_job_skills, jd_text)]
    for f_bytes, f_name in file_contents:
        tasks.append(asyncio.to_thread(parse_resume, f_bytes, f_name))
        
    # 3. Wait for all extractions simultaneously
    parallel_results = await asyncio.gather(*tasks)
    
    job_skills = parallel_results[0]
    resume_results = parallel_results[1:]
    
    ranked_candidates = []
    
    # Run matches
    for idx, (f_bytes, f_name) in enumerate(file_contents):
        extracted_data = resume_results[idx]
        
        if "error" in extracted_data:
            ranked_candidates.append({
                "name": f_name,
                "score": 0,
                "decision": "REJECTED",
                "tier": "Low Tier",
                "keyword_score": 0.0,
                "semantic_score": 0.0,
                "matched_skills": [],
                "missing_required": []
            })
            continue
            
        resume_skills = extracted_data.get("skills", [])
        
        # Match
        match_result = match_skills(resume_skills, job_skills, jd_text)
        score = match_result["score"]
        decision = classify_decision(score, t)
        tier = classify_tier(score, t)
        add_analysis_record(current_user["id"], f_name, score, decision, tier)
        
        ranked_candidates.append({
            "name": f_name,
            "score": score,
            "decision": decision,
            "tier": tier,
            "keyword_score": match_result.get("keyword_score", 0.0),
            "semantic_score": match_result.get("semantic_score", 0.0),
            "matched_skills": match_result.get("matched_skills", []),
            "matched_required": match_result.get("matched_required", []),
            "matched_optional": match_result.get("matched_optional", []),
            "missing_required": match_result.get("missing_required", []),
            "missing_optional": match_result.get("missing_optional", []),
        })
        
    ranked_candidates.sort(key=lambda x: x["score"], reverse=True)
    return {"ranked_candidates": ranked_candidates, "thresholds": t}

@app.post("/batch-match")
async def batch_match_endpoint(
    resumes: List[UploadFile] = File(...),
    job_description_files: List[UploadFile] = File(...),
    thresholds: Optional[str] = Form(None),
):
    jd_texts = []
    
    for file in job_description_files:
        try:
            content = await file.read()
            filename = file.filename
            ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
            extracted_text = ""
            
            if ext == "pdf":
                import pdfplumber
                import io
                text_parts = []
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    for page in pdf.pages:
                        t = page.extract_text()
                        if t:
                            text_parts.append(t)
                extracted_text = "\n".join(text_parts)
            elif ext in ("docx", "doc"):
                import docx
                import io
                doc = docx.Document(io.BytesIO(content))
                extracted_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            elif ext == "txt":
                try:
                    extracted_text = content.decode("utf-8", errors="replace")
                except Exception:
                    extracted_text = content.decode("latin-1", errors="replace")
                    
            if extracted_text:
                clean_lines = [line.strip() for line in extracted_text.split("\n") if line.strip()]
                final_text = "\n".join(clean_lines)
                if final_text:
                    jd_texts.append(final_text)
                    
        except Exception as e:
            print(f"Error processing {file.filename}: {e}")
            
    print("Total JDs:", len(jd_texts))

    if not jd_texts:
        return {"error": "Could not extract text from the provided files."}
    parsed_resumes = []
    for file in resumes:
        try:
            file_bytes = await file.read()
            result = parse_resume(file_bytes, file.filename)
            
            if "error" not in result:
                skills = result.get("skills", [])
                parsed_resumes.append({
                    "filename": file.filename,
                    "skills": skills
                })
        except Exception as e:
            print(f"Error parsing resume {file.filename}: {e}")
            continue

    if not parsed_resumes:
        return {"error": "Could not parse any of the provided resumes."}

    # Extract job skills from each JD upfront to avoid redundant LLM calls
    job_skills_list = []
    for text in jd_texts:
        job_skills = extract_job_skills(text)
        job_skills_list.append(job_skills)

    all_matches = []

    for resume in parsed_resumes:
        resume_skills = resume["skills"]
        filename = resume["filename"]

        matches = []

        for j in range(len(jd_texts)):
            jd_text = jd_texts[j]
            job_skills = job_skills_list[j]

            # Reusing existing matcher logic
            score_result = match_skills(resume_skills, job_skills, jd_text)

            # Extract score
            score = score_result.get("score", 0)

            matches.append({
                "job_id": j,
                "score": score
            })

        all_matches.append({
            "filename": filename,
            "matches": matches
        })

    # Parse dynamic frontend threshold settings
    t = DEFAULT_THRESHOLDS.copy()
    if thresholds:
        try:
            t.update(json.loads(thresholds))
        except (json.JSONDecodeError, TypeError):
            pass

    # Step 5 & 6: Cross-distribute candidate across ALL matched jobs natively
    response = []
    for item in all_matches:
        filename = item["filename"]
        for match in item["matches"]:
            job_id = match["job_id"]
            score = match["score"]
            
            job_name = f"Job {job_id + 1}"
            tier = classify_tier(score, t)

            response.append({
                "candidate": filename,
                "best_job": job_name,
                "score": score,
                "tier": tier
            })
        
    # Sort descending by score for better frontend layout
    response.sort(key=lambda x: x["score"], reverse=True)

    grouped = {}
    for result in response:
        job = result["best_job"]
        tier = result["tier"]

        if job not in grouped:
            grouped[job] = {
                "Top Tier": [],
                "Mid Tier": [],
                "Low Tier": []
            }
        
        grouped[job][tier].append(result)

    return {
        "grouped_results": grouped
    }



@app.post("/report")
async def report_endpoint(
    file: UploadFile = File(...),
    job_description: str = Form(...)
):
    file_bytes = await file.read()
    
    # Extract file and job skills concurrently
    extracted_data, job_skills = await asyncio.gather(
        asyncio.to_thread(parse_resume, file_bytes, file.filename),
        asyncio.to_thread(extract_job_skills, job_description)
    )
    
    resume_skills = extracted_data.get("skills", [])
    
    # 3. Match skills
    match_result = match_skills(resume_skills, job_skills, job_description)
    
    pdf_path = generate_pdf(match_result)
    
    return FileResponse(
        path=pdf_path,
        filename=f"Report_{file.filename}.pdf",
        media_type="application/pdf"
    )

@app.post("/extract-tables")
async def extract_tables_endpoint(file: UploadFile = File(...)):
    """Extract all tables from a PDF and return structured JSON."""
    file_bytes = await file.read()
    ext = file.filename.lower().split(".")[-1]
    if ext != "pdf":
        return {"error": "Only PDF files are supported for table extraction."}
    result = extract_tables_from_bytes(file_bytes)
    return result
